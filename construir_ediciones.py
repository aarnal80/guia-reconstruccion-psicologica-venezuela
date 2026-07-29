from __future__ import annotations

import html
import os
import re
import shutil
import zipfile
from pathlib import Path
from xml.sax.saxutils import escape

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from reportlab.lib.colors import HexColor, white
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT, TA_RIGHT
from reportlab.lib.pagesizes import inch
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfgen import canvas as pdfcanvas
from reportlab.lib.utils import ImageReader
from reportlab.platypus import (
    BaseDocTemplate,
    Frame,
    KeepTogether,
    ListFlowable,
    ListItem,
    PageBreak,
    PageTemplate,
    Paragraph,
    Spacer,
    Table,
    TableStyle,
)


ROOT = Path(__file__).resolve().parent
SOURCE = ROOT / "manuscrito_guia.md"
STEM = "Guia_reconstruccion_psicologica_Indira_Parra_Antonio_Arnal"
DOCX_PATH = ROOT / f"{STEM}.docx"
PDF_PATH = ROOT / f"{STEM}_KDP_6x9.pdf"
HTML_PATH = ROOT / f"{STEM}_web.html"
EPUB_PATH = ROOT / f"{STEM}.epub"
FRONT_COVER_PDF = ROOT / "Portada_provisional_frontal_6x9.pdf"
WRAP_COVER_PDF = ROOT / "Portada_provisional_KDP_90pag_crema.pdf"
COVER_ART = ROOT / "arte_portada_venezuela_luto_luminosa.png"

TITLE = "Guía de reconstrucción psicológica"
SUBTITLE = "Cuando todo se derrumba por dentro y por fuera"
AUTHORS = "Indira Lucía Parra y Antonio José Arnal Meinhardt"

# Named design override:
# Base preset: narrative_proposal.
# Override: KDP 6x9 narrative guide, no bleed, Georgia, 0.82" margins,
# quiet teal/navy palette, 11.4 pt body, generous reading spacing.
NAVY = "17324D"
TEAL = "2D6F73"
GOLD = "9A6B32"
MUTED = "66737C"
PALE = "EEF5F4"
INK = "1F2529"

TOC_ENTRIES = [
    ("Introducción", "Por qué nace esta guía", 7),
    ("Guía 1", "Comprender lo que sentimos es el primer paso para reconstruirnos", 11),
    ("Guía 2", "¿Qué necesita nuestro cerebro para comenzar a recuperarse?", 21),
    ("Guía 3", "¿Cómo atravesar el duelo sin dejar de vivir?", 32),
    ("Guía 4", "¿Cómo ayudamos a quien está sufriendo?", 45),
    ("Guía 5", "Niños, niñas y adolescentes", 54),
    ("Guía 6", "Cuando quien cuida también necesita ser cuidado", 63),
    ("Guía 7", "Reconstruirse: volver a habitar la vida", 70),
    ("Herramientas prácticas", "Hojas de trabajo y planes breves", 78),
    ("Palabras para seguir caminando", "", 83),
    ("Referencias", "", 85),
    ("Sobre los autores", "", 88),
]


def parse_markdown(text: str):
    lines = text.replace("\r\n", "\n").split("\n")
    blocks = []
    i = 0
    while i < len(lines):
        source_line = lines[i]
        raw = source_line.rstrip()
        if not raw.strip():
            i += 1
            continue
        if raw.strip() == "<!-- PAGEBREAK -->":
            blocks.append(("pagebreak", "", None))
            i += 1
            continue
        m = re.match(r"^(#{1,3})\s+(.*)$", raw)
        if m:
            blocks.append(("heading", m.group(2).strip(), len(m.group(1))))
            i += 1
            continue
        if raw.startswith("> "):
            parts = []
            while i < len(lines) and lines[i].startswith("> "):
                parts.append(lines[i][2:].strip())
                i += 1
            blocks.append(("quote", " ".join(parts), None))
            continue
        if re.match(r"^\s*[-*]\s+", raw):
            items = []
            while i < len(lines) and re.match(r"^\s*[-*]\s+", lines[i]):
                items.append(re.sub(r"^\s*[-*]\s+", "", lines[i]).strip())
                i += 1
            blocks.append(("ul", items, None))
            continue
        if re.match(r"^\s*\d+\.\s+", raw):
            items = []
            while i < len(lines) and re.match(r"^\s*\d+\.\s+", lines[i]):
                items.append(re.sub(r"^\s*\d+\.\s+", "", lines[i]).strip())
                i += 1
            blocks.append(("ol", items, None))
            continue
        if source_line.endswith("  "):
            blocks.append(("hardline", raw.strip(), None))
            i += 1
            continue
        parts = [raw.strip()]
        i += 1
        while i < len(lines):
            next_source_line = lines[i]
            nxt = next_source_line.rstrip()
            if (
                not nxt.strip()
                or nxt.strip() == "<!-- PAGEBREAK -->"
                or re.match(r"^(#{1,3})\s+", nxt)
                or nxt.startswith("> ")
                or re.match(r"^\s*[-*]\s+", nxt)
                or re.match(r"^\s*\d+\.\s+", nxt)
            ):
                break
            parts.append(nxt.strip())
            i += 1
            if next_source_line.endswith("  "):
                break
        blocks.append(("paragraph", " ".join(parts), None))
    return blocks


INLINE_RE = re.compile(r"(\*\*.*?\*\*|https?://[^\s<]+)")


def inline_segments(text: str):
    pos = 0
    for m in INLINE_RE.finditer(text):
        if m.start() > pos:
            yield "text", text[pos:m.start()]
        token = m.group(0)
        if token.startswith("**"):
            yield "bold", token[2:-2]
        else:
            yield "url", token
        pos = m.end()
    if pos < len(text):
        yield "text", text[pos:]


def reportlab_inline(text: str) -> str:
    out = []
    for kind, value in inline_segments(text):
        safe = escape(value)
        if kind == "bold":
            out.append(f"<b>{safe}</b>")
        elif kind == "url":
            out.append(f'<link href="{safe}" color="#2D6F73">{safe}</link>')
        else:
            out.append(safe)
    return "".join(out)


def add_docx_runs(paragraph, text: str, size=None, color=None, italic=False):
    for kind, value in inline_segments(text):
        run = paragraph.add_run(value)
        run.font.name = "Georgia"
        run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Georgia")
        run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Georgia")
        run.bold = kind == "bold"
        run.italic = italic
        if size:
            run.font.size = Pt(size)
        if color:
            run.font.color.rgb = RGBColor.from_string(color)


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tcMar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tcMar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char1, instr, fld_char2])


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def new_numbering_instance(doc: Document, abstract_num_id: int) -> int:
    numbering = doc.part.numbering_part.element
    existing = [
        int(node.get(qn("w:numId")))
        for node in numbering.findall(qn("w:num"))
        if node.get(qn("w:numId"))
    ]
    num_id = max(existing, default=0) + 1
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract = OxmlElement("w:abstractNumId")
    abstract.set(qn("w:val"), str(abstract_num_id))
    num.append(abstract)
    override = OxmlElement("w:lvlOverride")
    override.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:startOverride")
    start.set(qn("w:val"), "1")
    override.append(start)
    num.append(override)
    numbering.append(num)
    return num_id


def apply_numbering(paragraph, num_id: int):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num)


def set_table_geometry(table, widths_dxa):
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "0")
    tbl_ind.set(qn("w:type"), "dxa")
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = OxmlElement(f"w:{edge}")
        node.set(qn("w:val"), "nil")
        borders.append(node)
    tbl_pr.append(borders)
    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(widths_dxa[idx] / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[idx]))
            tc_w.set(qn("w:type"), "dxa")


def add_docx_toc(doc):
    table = doc.add_table(rows=0, cols=2)
    widths = [5680, 360]
    for label, title, page in TOC_ENTRIES:
        cells = table.add_row().cells
        set_cell_margins(cells[0], top=55, start=0, bottom=55, end=90)
        set_cell_margins(cells[1], top=55, start=0, bottom=55, end=0)
        p = cells[0].paragraphs[0]
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        label_run = p.add_run(label)
        label_run.font.name = "Georgia"
        label_run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Georgia")
        label_run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Georgia")
        label_run.font.size = Pt(10.2)
        label_run.bold = True
        label_run.font.color.rgb = RGBColor.from_string(TEAL)
        if title:
            title_run = p.add_run()
            title_run.add_break()
            title_run.add_text(title)
            title_run.font.name = "Georgia"
            title_run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Georgia")
            title_run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Georgia")
            title_run.font.size = Pt(9.1)
            title_run.font.color.rgb = RGBColor.from_string(MUTED)
        pp = cells[1].paragraphs[0]
        pp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        pp.paragraph_format.space_before = Pt(0)
        pp.paragraph_format.space_after = Pt(0)
        page_run = pp.add_run(str(page))
        page_run.font.name = "Georgia"
        page_run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Georgia")
        page_run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Georgia")
        page_run.font.size = Pt(10)
        page_run.bold = True
        page_run.font.color.rgb = RGBColor.from_string(GOLD)
    set_table_geometry(table, widths)
    return table


def configure_docx_styles(doc: Document):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Georgia"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Georgia")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Georgia")
    normal.font.size = Pt(11.4)
    normal.font.color.rgb = RGBColor.from_string(INK)
    language = OxmlElement("w:lang")
    language.set(qn("w:val"), "es-VE")
    language.set(qn("w:eastAsia"), "es-VE")
    normal._element.rPr.append(language)
    pf = normal.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.space_before = Pt(0)
    pf.space_after = Pt(8.5)
    pf.line_spacing = 1.38
    pf.widow_control = True

    heading_tokens = {
        "Title": (27, NAVY, 0, 10, WD_ALIGN_PARAGRAPH.CENTER),
        "Subtitle": (14, TEAL, 0, 18, WD_ALIGN_PARAGRAPH.CENTER),
        "Heading 1": (18, NAVY, 18, 10, WD_ALIGN_PARAGRAPH.LEFT),
        "Heading 2": (14, TEAL, 14, 7, WD_ALIGN_PARAGRAPH.LEFT),
        "Heading 3": (11.5, GOLD, 10, 5, WD_ALIGN_PARAGRAPH.LEFT),
    }
    if "Subtitle" not in styles:
        styles.add_style("Subtitle", WD_STYLE_TYPE.PARAGRAPH)
    for name, (size, color, before, after, align) in heading_tokens.items():
        s = styles[name]
        s.font.name = "Georgia"
        s._element.rPr.rFonts.set(qn("w:ascii"), "Georgia")
        s._element.rPr.rFonts.set(qn("w:hAnsi"), "Georgia")
        s.font.size = Pt(size)
        s.font.bold = name != "Subtitle"
        s.font.color.rgb = RGBColor.from_string(color)
        s.paragraph_format.space_before = Pt(before)
        s.paragraph_format.space_after = Pt(after)
        s.paragraph_format.alignment = align
        s.paragraph_format.keep_with_next = True
        s.paragraph_format.keep_together = True

    for name in ("List Bullet", "List Number"):
        s = styles[name]
        s.font.name = "Georgia"
        s._element.rPr.rFonts.set(qn("w:ascii"), "Georgia")
        s._element.rPr.rFonts.set(qn("w:hAnsi"), "Georgia")
        s.font.size = Pt(11.4)
        s.paragraph_format.left_indent = Inches(0.34)
        s.paragraph_format.first_line_indent = Inches(-0.17)
        s.paragraph_format.space_after = Pt(4)
        s.paragraph_format.line_spacing = 1.25

    if "Block Quote" not in styles:
        styles.add_style("Block Quote", WD_STYLE_TYPE.PARAGRAPH)
    quote = styles["Block Quote"]
    quote.font.name = "Georgia"
    quote._element.rPr.rFonts.set(qn("w:ascii"), "Georgia")
    quote._element.rPr.rFonts.set(qn("w:hAnsi"), "Georgia")
    quote.font.size = Pt(11)
    quote.font.italic = True
    quote.font.color.rgb = RGBColor.from_string(NAVY)
    quote.paragraph_format.left_indent = Inches(0.3)
    quote.paragraph_format.right_indent = Inches(0.2)
    quote.paragraph_format.space_before = Pt(7)

    if "Worksheet" not in styles:
        styles.add_style("Worksheet", WD_STYLE_TYPE.PARAGRAPH)
    worksheet = styles["Worksheet"]
    worksheet.font.name = "Georgia"
    worksheet._element.rPr.rFonts.set(qn("w:ascii"), "Georgia")
    worksheet._element.rPr.rFonts.set(qn("w:hAnsi"), "Georgia")
    worksheet.font.size = Pt(11.4)
    worksheet.font.color.rgb = RGBColor.from_string(INK)
    worksheet.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    worksheet.paragraph_format.space_before = Pt(0)
    worksheet.paragraph_format.space_after = Pt(8.5)
    worksheet.paragraph_format.line_spacing = 1.28
    worksheet.paragraph_format.widow_control = True
    quote.paragraph_format.space_after = Pt(9)
    quote.paragraph_format.line_spacing = 1.25

    if "Reference" not in styles:
        styles.add_style("Reference", WD_STYLE_TYPE.PARAGRAPH)
    ref = styles["Reference"]
    ref.font.name = "Georgia"
    ref._element.rPr.rFonts.set(qn("w:ascii"), "Georgia")
    ref._element.rPr.rFonts.set(qn("w:hAnsi"), "Georgia")
    ref.font.size = Pt(9.2)
    ref.font.color.rgb = RGBColor.from_string(INK)
    ref.paragraph_format.left_indent = Inches(0.16)
    ref.paragraph_format.first_line_indent = Inches(-0.16)
    ref.paragraph_format.space_after = Pt(6)
    ref.paragraph_format.line_spacing = 1.15


def build_docx(blocks):
    doc = Document()
    auto_hyphenation = OxmlElement("w:autoHyphenation")
    auto_hyphenation.set(qn("w:val"), "true")
    doc.settings.element.append(auto_hyphenation)
    section = doc.sections[0]
    section.page_width = Inches(6)
    section.page_height = Inches(9)
    section.top_margin = Inches(0.78)
    section.bottom_margin = Inches(0.78)
    section.left_margin = Inches(0.82)
    section.right_margin = Inches(0.82)
    section.header_distance = Inches(0.32)
    section.footer_distance = Inches(0.35)
    section.different_first_page_header_footer = True

    configure_docx_styles(doc)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = header.add_run("GUÍA DE RECONSTRUCCIÓN PSICOLÓGICA")
    run.font.name = "Georgia"
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor.from_string(MUTED)
    footer = section.footer.paragraphs[0]
    add_page_number(footer)
    for r in footer.runs:
        r.font.name = "Georgia"
        r.font.size = Pt(8)
        r.font.color.rgb = RGBColor.from_string(MUTED)

    in_references = False
    first_h1 = True
    guide_opener = False
    for kind, content, level in blocks:
        if kind == "pagebreak":
            doc.add_page_break()
            continue
        if kind == "heading":
            if level == 1:
                in_references = content == "Referencias"
                guide_opener = bool(re.match(r"^Guía \d+", content))
                if guide_opener:
                    spacer = doc.add_paragraph()
                    spacer.paragraph_format.space_before = Pt(0)
                    spacer.paragraph_format.space_after = Pt(0)
                    spacer.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
                    spacer.paragraph_format.line_spacing = Pt(216)
                    spacer_run = spacer.add_run("\u00a0")
                    spacer_run.font.size = Pt(1)
                p = doc.add_paragraph(style="Heading 1")
                p.paragraph_format.page_break_before = not first_h1 and not guide_opener
                if guide_opener:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p.paragraph_format.space_before = Pt(0)
                    p.paragraph_format.space_after = Pt(18)
                first_h1 = False
            elif level == 2:
                p = doc.add_paragraph(style="Heading 2")
            else:
                p = doc.add_paragraph(style="Heading 3")
            add_docx_runs(p, content)
            if level == 1 and guide_opener:
                for run in p.runs:
                    run.font.size = Pt(23)
            if level == 2 and guide_opener:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.size = Pt(16)
                doc.add_page_break()
                guide_opener = False
            continue
        if kind == "quote":
            p = doc.add_paragraph(style="Block Quote")
            add_docx_runs(p, content, italic=True)
            continue
        if kind in ("ul", "ol"):
            style = "List Bullet" if kind == "ul" else "List Number"
            num_id = new_numbering_instance(doc, 7) if kind == "ol" else None
            for item in content:
                p = doc.add_paragraph(style=style)
                if num_id is not None:
                    apply_numbering(p, num_id)
                add_docx_runs(p, item)
            continue
        if content == "[[TOC_STATIC]]":
            add_docx_toc(doc)
            continue
        style = (
            "Reference"
            if in_references
            else "Worksheet"
            if kind == "hardline" or re.search(r"_{3,}", content)
            else "Normal"
        )
        p = doc.add_paragraph(style=style)
        add_docx_runs(p, content)

    props = doc.core_properties
    props.title = TITLE
    props.subject = SUBTITLE
    props.author = AUTHORS
    props.keywords = "duelo, terremoto, Venezuela, primeros auxilios psicológicos, reconstrucción"
    props.comments = "Edición 6 x 9 pulgadas sin sangrado."
    doc.save(DOCX_PATH)


class BookDocTemplate(BaseDocTemplate):
    def __init__(self, filename, **kwargs):
        super().__init__(filename, **kwargs)
        frame = Frame(
            self.leftMargin,
            self.bottomMargin,
            self.width,
            self.height,
            id="body",
            leftPadding=0,
            rightPadding=0,
            topPadding=0,
            bottomPadding=0,
        )
        self.addPageTemplates(PageTemplate(id="book", frames=[frame], onPage=self.draw_page))
        self._bookmark_counter = 0

    def draw_page(self, canvas, doc):
        page = canvas.getPageNumber()
        if page > 1:
            canvas.saveState()
            canvas.setFont("Georgia", 7.5)
            canvas.setFillColor(HexColor(f"#{MUTED}"))
            canvas.drawCentredString(3 * inch, 8.55 * inch, "GUÍA DE RECONSTRUCCIÓN PSICOLÓGICA")
            canvas.drawCentredString(3 * inch, 0.35 * inch, str(page))
            canvas.restoreState()

    def afterFlowable(self, flowable):
        if isinstance(flowable, Paragraph) and flowable.style.name in (
            "H1",
            "H2",
            "GuideH1",
            "GuideH2",
        ):
            self._bookmark_counter += 1
            key = f"h{self._bookmark_counter}"
            text = flowable.getPlainText()
            self.canv.bookmarkPage(key)
            level = 0 if flowable.style.name in ("H1", "GuideH1") else 1
            self.canv.addOutlineEntry(text, key, level=level, closed=False)


def register_pdf_fonts():
    fonts = Path("C:/Windows/Fonts")
    pdfmetrics.registerFont(TTFont("Georgia", fonts / "georgia.ttf"))
    pdfmetrics.registerFont(TTFont("Georgia-Bold", fonts / "georgiab.ttf"))
    pdfmetrics.registerFont(TTFont("Georgia-Italic", fonts / "georgiai.ttf"))
    pdfmetrics.registerFont(TTFont("Georgia-BoldItalic", fonts / "georgiaz.ttf"))
    pdfmetrics.registerFontFamily(
        "Georgia",
        normal="Georgia",
        bold="Georgia-Bold",
        italic="Georgia-Italic",
        boldItalic="Georgia-BoldItalic",
    )


def pdf_styles():
    return {
        "Body": ParagraphStyle(
            "Body",
            fontName="Georgia",
            fontSize=11.4,
            leading=17.8,
            textColor=HexColor(f"#{INK}"),
            alignment=TA_JUSTIFY,
            spaceAfter=8.5,
            allowWidows=0,
            allowOrphans=0,
            splitLongWords=True,
            hyphenationLang="es",
            embeddedHyphenation=1,
        ),
        "H1": ParagraphStyle(
            "H1",
            fontName="Georgia-Bold",
            fontSize=18,
            leading=22,
            textColor=HexColor(f"#{NAVY}"),
            spaceBefore=10,
            spaceAfter=13,
            keepWithNext=True,
        ),
        "H2": ParagraphStyle(
            "H2",
            fontName="Georgia-Bold",
            fontSize=14,
            leading=18,
            textColor=HexColor(f"#{TEAL}"),
            spaceBefore=11,
            spaceAfter=7,
            keepWithNext=True,
        ),
        "H3": ParagraphStyle(
            "H3",
            fontName="Georgia-Bold",
            fontSize=11.5,
            leading=15,
            textColor=HexColor(f"#{GOLD}"),
            spaceBefore=9,
            spaceAfter=5,
            keepWithNext=True,
        ),
        "GuideH1": ParagraphStyle(
            "GuideH1",
            fontName="Georgia-Bold",
            fontSize=23,
            leading=28,
            textColor=HexColor(f"#{NAVY}"),
            alignment=TA_CENTER,
            spaceBefore=0,
            spaceAfter=18,
            keepWithNext=True,
        ),
        "GuideH2": ParagraphStyle(
            "GuideH2",
            fontName="Georgia-Bold",
            fontSize=16,
            leading=22,
            textColor=HexColor(f"#{TEAL}"),
            alignment=TA_CENTER,
            spaceBefore=0,
            spaceAfter=0,
            keepWithNext=True,
        ),
        "Worksheet": ParagraphStyle(
            "Worksheet",
            fontName="Georgia",
            fontSize=11.4,
            leading=16.4,
            textColor=HexColor(f"#{INK}"),
            alignment=TA_LEFT,
            spaceAfter=8.5,
            allowWidows=0,
            allowOrphans=0,
            splitLongWords=True,
            hyphenationLang="es",
            embeddedHyphenation=1,
        ),
        "Quote": ParagraphStyle(
            "Quote",
            fontName="Georgia-Italic",
            fontSize=11,
            leading=15.5,
            textColor=HexColor(f"#{NAVY}"),
            leftIndent=16,
            rightIndent=11,
            borderColor=HexColor(f"#{TEAL}"),
            borderWidth=1.2,
            borderPadding=(7, 9, 7, 10),
            backColor=HexColor(f"#{PALE}"),
            spaceBefore=6,
            spaceAfter=10,
        ),
        "List": ParagraphStyle(
            "List",
            fontName="Georgia",
            fontSize=11.3,
            leading=17.1,
            textColor=HexColor(f"#{INK}"),
            leftIndent=0,
            spaceAfter=3.5,
        ),
        "Reference": ParagraphStyle(
            "Reference",
            fontName="Georgia",
            fontSize=8.9,
            leading=11.8,
            textColor=HexColor(f"#{INK}"),
            leftIndent=10,
            firstLineIndent=-10,
            spaceAfter=5,
            alignment=TA_LEFT,
            splitLongWords=True,
            hyphenationLang="es",
            embeddedHyphenation=1,
        ),
    }


def pdf_toc_table():
    left_style = ParagraphStyle(
        "TocLeft",
        fontName="Georgia",
        fontSize=9.1,
        leading=12.2,
        textColor=HexColor(f"#{MUTED}"),
        spaceAfter=0,
    )
    page_style = ParagraphStyle(
        "TocPage",
        fontName="Georgia-Bold",
        fontSize=10,
        leading=12,
        textColor=HexColor(f"#{GOLD}"),
        alignment=TA_RIGHT,
        spaceAfter=0,
    )
    rows = []
    for label, title, page in TOC_ENTRIES:
        left = f'<font color="#{TEAL}"><b>{escape(label)}</b></font>'
        if title:
            left += f'<br/><font color="#{MUTED}" size="9">{escape(title)}</font>'
        rows.append([Paragraph(left, left_style), Paragraph(str(page), page_style)])
    table = Table(rows, colWidths=[3.86 * inch, 0.50 * inch], hAlign="LEFT")
    table.setStyle(
        TableStyle(
            [
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("LEFTPADDING", (0, 0), (-1, -1), 0),
                ("RIGHTPADDING", (0, 0), (0, -1), 7),
                ("RIGHTPADDING", (1, 0), (1, -1), 0),
                ("TOPPADDING", (0, 0), (-1, -1), 1.5),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 4.5),
                ("LINEBELOW", (0, 0), (-1, -2), 0.25, HexColor("#D8E3E3")),
            ]
        )
    )
    return table


def build_pdf(blocks):
    register_pdf_fonts()
    styles = pdf_styles()
    doc = BookDocTemplate(
        str(PDF_PATH),
        pagesize=(6 * inch, 9 * inch),
        leftMargin=0.82 * inch,
        rightMargin=0.82 * inch,
        topMargin=0.78 * inch,
        bottomMargin=0.78 * inch,
        title=TITLE,
        author=AUTHORS,
        subject=SUBTITLE,
    )
    story = []
    in_references = False
    guide_opener = False
    previous_pagebreak = False
    for kind, content, level in blocks:
        if kind == "pagebreak":
            if not previous_pagebreak:
                story.append(PageBreak())
            previous_pagebreak = True
            continue
        previous_pagebreak = False
        if kind == "heading":
            if level == 1:
                in_references = content == "Referencias"
                guide_opener = bool(re.match(r"^Guía \d+", content))
                if guide_opener:
                    story.append(Spacer(1, 3.0 * inch))
                    story.append(
                        Paragraph(reportlab_inline(content), styles["GuideH1"])
                    )
                else:
                    story.append(Paragraph(reportlab_inline(content), styles["H1"]))
                    story.append(Spacer(1, 4))
            elif level == 2:
                if guide_opener:
                    story.append(
                        Paragraph(reportlab_inline(content), styles["GuideH2"])
                    )
                    story.append(PageBreak())
                    guide_opener = False
                else:
                    story.append(Paragraph(reportlab_inline(content), styles["H2"]))
            else:
                story.append(Paragraph(reportlab_inline(content), styles["H3"]))
            continue
        if kind == "quote":
            story.append(Paragraph(reportlab_inline(content), styles["Quote"]))
            continue
        if kind in ("ul", "ol"):
            items = [
                ListItem(
                    Paragraph(reportlab_inline(item), styles["List"]),
                    leftIndent=10,
                    value=(idx + 1 if kind == "ol" else None),
                )
                for idx, item in enumerate(content)
            ]
            story.append(
                ListFlowable(
                    items,
                    bulletType="bullet" if kind == "ul" else "1",
                    start="circle" if kind == "ul" else "1",
                    leftIndent=18,
                    bulletFontName="Georgia",
                    bulletFontSize=9,
                    bulletOffsetY=1,
                    spaceAfter=7,
                )
            )
            continue
        if content == "[[TOC_STATIC]]":
            story.append(pdf_toc_table())
            continue
        style = (
            styles["Reference"]
            if in_references
            else styles["Worksheet"]
            if kind == "hardline" or re.search(r"_{3,}", content)
            else styles["Body"]
        )
        story.append(Paragraph(reportlab_inline(content), style))
    doc.build(story)


def html_inline(text: str) -> str:
    escaped = html.escape(text)
    escaped = re.sub(r"\*\*(.*?)\*\*", r"<strong>\1</strong>", escaped)
    escaped = re.sub(
        r"(https?://[^\s<]+)",
        r'<a href="\1">\1</a>',
        escaped,
    )
    return escaped


def static_toc_html():
    return '<div class="editorial-toc">' + "".join(
        (
            '<div class="toc-entry">'
            f'<div><strong>{html.escape(label)}</strong>'
            + (f"<span>{html.escape(title)}</span>" if title else "")
            + "</div>"
            f'<b class="toc-page">{page}</b>'
            "</div>"
        )
        for label, title, page in TOC_ENTRIES
    ) + "</div>"


def build_html(blocks):
    body = []
    toc = []
    slug_counts = {}
    for kind, content, level in blocks:
        if kind == "pagebreak":
            body.append('<hr class="pagebreak">')
        elif kind == "heading":
            base = re.sub(r"[^a-z0-9]+", "-", content.lower()
                          .replace("á", "a").replace("é", "e").replace("í", "i")
                          .replace("ó", "o").replace("ú", "u").replace("ñ", "n")).strip("-")
            slug_counts[base] = slug_counts.get(base, 0) + 1
            slug = base if slug_counts[base] == 1 else f"{base}-{slug_counts[base]}"
            body.append(f'<h{level} id="{slug}">{html_inline(content)}</h{level}>')
            if level <= 2 and content not in (TITLE.upper(),):
                toc.append((level, content, slug))
        elif kind == "quote":
            body.append(f"<blockquote>{html_inline(content)}</blockquote>")
        elif kind in ("ul", "ol"):
            tag = kind
            items = "".join(f"<li>{html_inline(x)}</li>" for x in content)
            body.append(f"<{tag}>{items}</{tag}>")
        elif content == "[[TOC_STATIC]]":
            body.append(static_toc_html())
        else:
            body.append(f"<p>{html_inline(content)}</p>")
    toc_html = "".join(
        f'<li class="l{level}"><a href="#{slug}">{html.escape(text)}</a></li>'
        for level, text, slug in toc
    )
    document = f"""<!doctype html>
<html lang="es">
<head>
<meta charset="utf-8">
<meta name="viewport" content="width=device-width,initial-scale=1">
<title>{TITLE} — {SUBTITLE}</title>
<meta name="author" content="{AUTHORS}">
<meta name="description" content="Primeros auxilios psicológicos ampliados para personas, familias y comunidades afectadas por catástrofes.">
<style>
:root{{--ink:#1f2529;--navy:#17324d;--teal:#2d6f73;--gold:#9a6b32;--pale:#eef5f4;--paper:#fffdf9}}
*{{box-sizing:border-box}} body{{margin:0;background:#edf1f2;color:var(--ink);font-family:Georgia,serif;line-height:1.65}}
.layout{{display:grid;grid-template-columns:280px minmax(0,760px);gap:34px;max-width:1120px;margin:auto;padding:28px}}
nav{{position:sticky;top:20px;align-self:start;max-height:calc(100vh - 40px);overflow:auto;background:white;padding:20px;border-radius:12px;box-shadow:0 8px 30px #17324d18}}
nav h2{{font-size:1rem;margin-top:0}} nav ul{{padding-left:0;list-style:none}} nav li{{margin:.35rem 0}} nav .l2{{padding-left:14px;font-size:.9rem}}
nav a{{color:var(--teal);text-decoration:none}} main{{background:var(--paper);padding:56px 68px;border-radius:12px;box-shadow:0 8px 30px #17324d18}}
h1{{font-size:2rem;color:var(--navy);line-height:1.2;margin:3rem 0 1rem}} h1:first-child{{margin-top:0}}
h2{{font-size:1.45rem;color:var(--teal);line-height:1.3;margin:2rem 0 .7rem}} h3{{font-size:1.05rem;color:var(--gold);margin:1.5rem 0 .4rem}}
p{{margin:.65rem 0;text-align:justify;hyphens:auto}} li{{margin:.35rem 0}} blockquote{{margin:1.3rem 0;padding:1rem 1.2rem;background:var(--pale);border-left:4px solid var(--teal);color:var(--navy);font-style:italic}}
.editorial-toc{{margin:1rem 0 2rem}}.toc-entry{{display:grid;grid-template-columns:1fr auto;gap:1rem;align-items:center;padding:.6rem 0;border-bottom:1px solid #d8e3e3}}
.toc-entry strong{{display:block;color:var(--teal);font-size:.96rem}}.toc-entry span{{display:block;color:var(--muted);font-size:.9rem;line-height:1.35}}.toc-page{{color:var(--gold);font-size:.95rem}}
.pagebreak{{border:0;border-top:1px solid #d9e2e1;margin:3rem 0}} a{{color:var(--teal)}} strong{{color:#132f3e}}
@media(max-width:850px){{.layout{{display:block;padding:0}}nav{{position:relative;top:0;max-height:none;border-radius:0}}main{{padding:34px 22px;border-radius:0}}}}
@media print{{body{{background:white}}nav{{display:none}}.layout{{display:block;padding:0}}main{{box-shadow:none;padding:0}}.pagebreak{{page-break-after:always;border:0}}}}
</style>
</head>
<body>
<div class="layout">
<nav aria-label="Contenido"><h2>Contenido</h2><ul>{toc_html}</ul></nav>
<main>{''.join(body)}</main>
</div>
</body>
</html>"""
    HTML_PATH.write_text(document, encoding="utf-8")


def build_epub(blocks):
    temp = ROOT / "_epub_build"
    if temp.exists():
        shutil.rmtree(temp)
    (temp / "META-INF").mkdir(parents=True)
    (temp / "OEBPS").mkdir(parents=True)
    (temp / "mimetype").write_text("application/epub+zip", encoding="ascii")
    (temp / "META-INF" / "container.xml").write_text(
        """<?xml version="1.0"?>
<container version="1.0" xmlns="urn:oasis:names:tc:opendocument:xmlns:container">
<rootfiles><rootfile full-path="OEBPS/content.opf" media-type="application/oebps-package+xml"/></rootfiles>
</container>""",
        encoding="utf-8",
    )
    css = """body{font-family:serif;line-height:1.5;color:#1f2529}p{text-align:justify;hyphens:auto}h1{color:#17324d}h2{color:#2d6f73}h3{color:#9a6b32}blockquote{border-left:3px solid #2d6f73;padding-left:1em;font-style:italic}.editorial-toc{margin:1em 0 2em}.toc-entry{display:flex;justify-content:space-between;gap:1em;padding:.5em 0;border-bottom:1px solid #d8e3e3}.toc-entry strong,.toc-entry span{display:block}.toc-entry strong{color:#2d6f73}.toc-entry span{color:#66737c;font-size:.9em}.toc-page{color:#9a6b32}.pagebreak{page-break-after:always}"""
    (temp / "OEBPS" / "style.css").write_text(css, encoding="utf-8")
    body = []
    nav = []
    count = 0
    for kind, content, level in blocks:
        if kind == "pagebreak":
            body.append('<div class="pagebreak"></div>')
        elif kind == "heading":
            count += 1
            slug = f"s{count}"
            body.append(f'<h{level} id="{slug}">{html_inline(content)}</h{level}>')
            if level <= 2:
                nav.append((level, content, slug))
        elif kind == "quote":
            body.append(f"<blockquote>{html_inline(content)}</blockquote>")
        elif kind in ("ul", "ol"):
            items = "".join(f"<li>{html_inline(x)}</li>" for x in content)
            body.append(f"<{kind}>{items}</{kind}>")
        elif content == "[[TOC_STATIC]]":
            body.append(static_toc_html())
        else:
            body.append(f"<p>{html_inline(content)}</p>")
    xhtml = f"""<?xml version="1.0" encoding="utf-8"?>
<!DOCTYPE html>
<html xmlns="http://www.w3.org/1999/xhtml" lang="es"><head><title>{TITLE}</title><link rel="stylesheet" href="style.css" type="text/css"/></head><body>{''.join(body)}</body></html>"""
    (temp / "OEBPS" / "book.xhtml").write_text(xhtml, encoding="utf-8")
    nav_items = "".join(
        f'<li{" class=\"sub\"" if level == 2 else ""}><a href="book.xhtml#{slug}">{html.escape(text)}</a></li>'
        for level, text, slug in nav
    )
    nav_xhtml = f"""<?xml version="1.0" encoding="utf-8"?>
<!DOCTYPE html>
<html xmlns="http://www.w3.org/1999/xhtml" xmlns:epub="http://www.idpf.org/2007/ops" lang="es"><head><title>Contenido</title></head><body><nav epub:type="toc"><h1>Contenido</h1><ol>{nav_items}</ol></nav></body></html>"""
    (temp / "OEBPS" / "nav.xhtml").write_text(nav_xhtml, encoding="utf-8")
    opf = f"""<?xml version="1.0" encoding="utf-8"?>
<package version="3.0" unique-identifier="pub-id" xmlns="http://www.idpf.org/2007/opf">
<metadata xmlns:dc="http://purl.org/dc/elements/1.1/">
<dc:identifier id="pub-id">urn:uuid:guia-reconstruccion-psicologica-2026</dc:identifier>
<dc:title>{TITLE}</dc:title><dc:creator>{AUTHORS}</dc:creator><dc:language>es</dc:language>
<meta property="dcterms:modified">2026-07-29T00:00:00Z</meta>
</metadata>
<manifest>
<item id="book" href="book.xhtml" media-type="application/xhtml+xml"/>
<item id="nav" href="nav.xhtml" media-type="application/xhtml+xml" properties="nav"/>
<item id="css" href="style.css" media-type="text/css"/>
</manifest>
<spine><itemref idref="book"/></spine>
</package>"""
    (temp / "OEBPS" / "content.opf").write_text(opf, encoding="utf-8")

    if EPUB_PATH.exists():
        EPUB_PATH.unlink()
    with zipfile.ZipFile(EPUB_PATH, "w") as zf:
        zf.write(temp / "mimetype", "mimetype", compress_type=zipfile.ZIP_STORED)
        for path in sorted(temp.rglob("*")):
            if path.is_file() and path.name != "mimetype":
                zf.write(path, path.relative_to(temp).as_posix(), compress_type=zipfile.ZIP_DEFLATED)
    shutil.rmtree(temp)


def draw_front_panel(canvas, x, y, width, height):
    navy = HexColor(f"#{NAVY}")
    teal = HexColor("#56A3A3")
    gold = HexColor(f"#{GOLD}")
    if COVER_ART.exists():
        canvas.drawImage(
            ImageReader(str(COVER_ART)),
            x,
            y,
            width=width,
            height=height,
            preserveAspectRatio=True,
            anchor="c",
            mask="auto",
        )
    else:
        canvas.setFillColor(navy)
        canvas.rect(x, y, width, height, stroke=0, fill=1)

    margin = 0.58 * inch
    canvas.setFillColor(gold)
    canvas.rect(x + margin, y + height - 0.72 * inch, 0.72 * inch, 0.045 * inch, stroke=0, fill=1)
    canvas.setFont("Georgia-Bold", 7.6)
    canvas.setFillColor(teal)
    canvas.drawString(x + margin, y + height - 0.98 * inch, "PRIMEROS AUXILIOS PSICOLÓGICOS AMPLIADOS")

    title_x = x + margin
    title_y = y + height - 1.70 * inch
    canvas.setFillColor(navy)
    canvas.setFont("Georgia-Bold", 27)
    canvas.drawString(title_x, title_y, "GUÍA DE")
    canvas.setFont("Georgia-Bold", 26)
    canvas.drawString(title_x, title_y - 0.48 * inch, "RECONSTRUCCIÓN")
    canvas.drawString(title_x, title_y - 0.96 * inch, "PSICOLÓGICA")

    canvas.setStrokeColor(teal)
    canvas.setLineWidth(1.1)
    canvas.line(title_x, title_y - 1.26 * inch, x + width - margin, title_y - 1.26 * inch)

    subtitle_style = ParagraphStyle(
        "CoverSubtitle",
        fontName="Georgia-Italic",
        fontSize=13.2,
        leading=17,
        textColor=navy,
    )
    subtitle = Paragraph("Cuando todo se derrumba<br/>por dentro y por fuera", subtitle_style)
    subtitle.wrapOn(canvas, 3.55 * inch, 0.8 * inch)
    subtitle.drawOn(canvas, title_x, title_y - 2.05 * inch)

    canvas.setFillColor(navy)
    canvas.setFont("Georgia-Bold", 10.2)
    canvas.drawString(title_x, y + 0.94 * inch, "INDIRA LUCÍA PARRA")
    canvas.setFont("Georgia", 9.6)
    canvas.drawString(title_x, y + 0.69 * inch, "ANTONIO JOSÉ ARNAL MEINHARDT")
    canvas.setFont("Georgia", 6.8)
    canvas.setFillColor(teal)
    canvas.drawRightString(x + width - margin, y + 0.43 * inch, "EDICIÓN PROVISIONAL · 2026")


def build_covers():
    register_pdf_fonts()

    front = pdfcanvas.Canvas(
        str(FRONT_COVER_PDF),
        pagesize=(6 * inch, 9 * inch),
        pageCompression=1,
    )
    front.setTitle(f"Portada provisional — {TITLE}")
    front.setAuthor(AUTHORS)
    draw_front_panel(front, 0, 0, 6 * inch, 9 * inch)
    front.showPage()
    front.save()

    page_count = 90
    bleed = 0.125 * inch
    trim_w = 6 * inch
    trim_h = 9 * inch
    spine_w = page_count * 0.0025 * inch
    wrap_w = 2 * bleed + 2 * trim_w + spine_w
    wrap_h = 2 * bleed + trim_h
    cover = pdfcanvas.Canvas(str(WRAP_COVER_PDF), pagesize=(wrap_w, wrap_h), pageCompression=1)
    cover.setTitle(f"Cubierta provisional KDP — {TITLE}")
    cover.setAuthor(AUTHORS)
    cover.setFillColor(HexColor("#D8EAF0"))
    cover.rect(0, 0, wrap_w, wrap_h, stroke=0, fill=1)

    back_x = bleed
    panel_y = bleed
    spine_x = back_x + trim_w
    front_x = spine_x + spine_w
    draw_front_panel(cover, front_x, panel_y, trim_w, trim_h)

    cover.setFillColor(HexColor("#EAF3F4"))
    cover.rect(back_x, panel_y, trim_w, trim_h, stroke=0, fill=1)
    cover.setFillColor(HexColor(f"#{GOLD}"))
    cover.rect(back_x + 0.58 * inch, panel_y + trim_h - 0.72 * inch, 0.72 * inch, 0.045 * inch, stroke=0, fill=1)

    cover.setFillColor(HexColor(f"#{NAVY}"))
    cover.setFont("Georgia-Bold", 16)
    cover.drawString(back_x + 0.58 * inch, panel_y + trim_h - 1.24 * inch, "Después de la catástrofe")

    blurb_style = ParagraphStyle(
        "BackBlurb",
        fontName="Georgia",
        fontSize=10.4,
        leading=15,
        textColor=HexColor(f"#{INK}"),
        spaceAfter=8,
    )
    blurb = Paragraph(
        "Cuando todo cambia de golpe, el miedo, la confusión, el insomnio, la rabia "
        "o el silencio pueden sentirse insoportables. Esta guía explica esas reacciones "
        "con palabras sencillas y ofrece pasos concretos para recuperar seguridad, "
        "acompañar el duelo y volver a habitar la vida.",
        blurb_style,
    )
    blurb.wrapOn(cover, 4.75 * inch, 2.0 * inch)
    blurb.drawOn(cover, back_x + 0.58 * inch, panel_y + trim_h - 3.25 * inch)

    bullet_style = ParagraphStyle(
        "BackBullets",
        fontName="Georgia",
        fontSize=9.4,
        leading=13.4,
        textColor=HexColor(f"#{NAVY}"),
        leftIndent=11,
        firstLineIndent=-11,
        bulletIndent=0,
    )
    bullets = [
        "Comprender qué ocurre en el cuerpo y la mente.",
        "Atravesar el duelo sin dejar de vivir.",
        "Acompañar a adultos, niños y adolescentes.",
        "Cuidar a quienes ayudan y sostienen a otros.",
        "Usar planes y ejercicios breves en momentos difíciles.",
    ]
    y_cursor = panel_y + trim_h - 3.72 * inch
    for item in bullets:
        paragraph = Paragraph(f"•&nbsp;&nbsp;{item}", bullet_style)
        _, paragraph_h = paragraph.wrap(4.65 * inch, 0.55 * inch)
        paragraph.drawOn(cover, back_x + 0.62 * inch, y_cursor - paragraph_h)
        y_cursor -= paragraph_h + 5

    author_style = ParagraphStyle(
        "BackAuthors",
        fontName="Georgia-Italic",
        fontSize=8.4,
        leading=11.5,
        textColor=HexColor("#466B77"),
    )
    author_text = Paragraph(
        "<b>Indira Lucía Parra</b>, médica psiquiatra, y <b>Antonio José Arnal "
        "Meinhardt</b> reúnen experiencia profesional y humana en duelo, atención a "
        "víctimas y respuesta ante emergencias en Venezuela.",
        author_style,
    )
    author_text.wrapOn(cover, 4.70 * inch, 1.1 * inch)
    author_text.drawOn(cover, back_x + 0.58 * inch, panel_y + 1.76 * inch)

    # Reserva segura para que KDP coloque el código de barras.
    barcode_x = back_x + trim_w - 2.55 * inch
    barcode_y = panel_y + 0.42 * inch
    cover.setFillColor(white)
    cover.roundRect(barcode_x, barcode_y, 2.0 * inch, 1.2 * inch, 3, stroke=0, fill=1)
    cover.setFillColor(HexColor("#A6A6A6"))
    cover.setFont("Georgia", 6.2)
    cover.drawCentredString(barcode_x + inch, barcode_y + 0.56 * inch, "ZONA DE CÓDIGO DE BARRAS")
    cover.setFillColor(HexColor("#56A3A3"))
    cover.setFont("Georgia", 6.5)
    cover.drawString(back_x + 0.58 * inch, panel_y + 0.48 * inch, "PORTADA PROVISIONAL")

    cover.setFillColor(HexColor("#4E8E9D"))
    cover.rect(spine_x, panel_y, spine_w, trim_h, stroke=0, fill=1)
    cover.saveState()
    cover.translate(spine_x + spine_w / 2, panel_y + 0.55 * inch)
    cover.rotate(90)
    cover.setFillColor(HexColor("#F4F0E8"))
    cover.setFont("Georgia-Bold", 7.5)
    cover.drawCentredString(3.95 * inch, -2.6, "GUÍA DE RECONSTRUCCIÓN PSICOLÓGICA")
    cover.restoreState()

    cover.showPage()
    cover.save()


def main():
    text = SOURCE.read_text(encoding="utf-8")
    blocks = parse_markdown(text)
    build_docx(blocks)
    build_pdf(blocks)
    build_html(blocks)
    build_epub(blocks)
    build_covers()
    print(f"Built {DOCX_PATH.name}")
    print(f"Built {PDF_PATH.name}")
    print(f"Built {HTML_PATH.name}")
    print(f"Built {EPUB_PATH.name}")
    print(f"Built {FRONT_COVER_PDF.name}")
    print(f"Built {WRAP_COVER_PDF.name}")


if __name__ == "__main__":
    main()
